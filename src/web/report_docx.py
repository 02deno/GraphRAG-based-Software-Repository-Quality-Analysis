"""Build a single Word (.docx) report bundling text artifacts and chart PNGs for one run."""

from __future__ import annotations

import json
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Tuple

from docx import Document
from docx.shared import Inches

from src.web.services.analysis_service import collect_visual_gallery_entries


def _append_preformatted(document: Document, text: str) -> None:
    """Append *text* line-by-line (readable in Word; long lines wrap naturally)."""
    body = text if text else "(empty)"
    for line in body.splitlines():
        document.add_paragraph(line)


def build_analysis_docx_bytes(run_dir: Path) -> Tuple[bytes, str]:
    """Assemble a .docx containing pipeline log, analysis, visual summary, stats, and PNGs.

    Args:
        run_dir: Resolved ``results/web_analysis_*`` directory.

    Returns:
        ``(docx_bytes, suggested_download_filename)``.
    """
    document = Document()
    stamp = datetime.now(tz=timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    document.add_heading("GraphRAG analysis report", level=0)
    document.add_paragraph(f"Run folder: {run_dir.name}")
    document.add_paragraph(f"Document generated (UTC): {stamp}")

    graph_path = run_dir / "graph.json"
    if graph_path.exists():
        document.add_heading("Graph overview", level=1)
        data = json.loads(graph_path.read_text(encoding="utf-8"))
        nodes = data.get("nodes") or []
        edges = data.get("edges") or []
        document.add_paragraph(f"Schema version: {data.get('schema_version', '')}")
        document.add_paragraph(f"Total nodes: {len(nodes)} · Total edges: {len(edges)}")
        document.add_paragraph(
            "Node types (implemented): "
            + ", ".join(data.get("implemented_node_types") or [])[:2000]
        )
        document.add_paragraph(
            "Edge types (implemented): "
            + ", ".join(data.get("implemented_edge_types") or [])[:2000]
        )
        document.add_paragraph(
            "Full ``graph.json`` is available as a separate download from the results page."
        )

    pipeline_path = run_dir / "pipeline.txt"
    document.add_heading("Pipeline log", level=1)
    if pipeline_path.exists():
        _append_preformatted(document, pipeline_path.read_text(encoding="utf-8"))
    else:
        document.add_paragraph("(No pipeline.txt for this run — older export.)")

    analysis_path = run_dir / "analysis.txt"
    document.add_heading("Degree / metrics report (analysis.txt)", level=1)
    if analysis_path.exists():
        _append_preformatted(document, analysis_path.read_text(encoding="utf-8"))
    else:
        document.add_paragraph("(analysis.txt missing)")

    vis_path = run_dir / "visual_summary.txt"
    document.add_heading("Visual summary (text)", level=1)
    if vis_path.exists():
        _append_preformatted(document, vis_path.read_text(encoding="utf-8"))
    else:
        document.add_paragraph("(visual_summary.txt missing)")

    gallery = collect_visual_gallery_entries(run_dir)
    if gallery:
        document.add_heading("Visualization charts (embedded PNGs)", level=1)
        for item in gallery:
            png = run_dir / "visuals" / item["name"]
            if not png.is_file():
                continue
            title = item.get("title") or item["name"]
            document.add_heading(title, level=2)
            try:
                document.add_picture(str(png), width=Inches(6.2))
            except OSError:
                document.add_paragraph(f"(Could not embed image: {item['name']})")

    buf = BytesIO()
    document.save(buf)
    fname = f"{run_dir.name}_full_report.docx"
    return buf.getvalue(), fname
